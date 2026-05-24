"""
Resume file parsing — PDF and DOCX → plain text.
Uses pdfplumber (layout-aware) and python-docx.
"""
import io
import pdfplumber
from docx import Document


def extract_text_from_pdf(content: bytes) -> str:
    text_parts = []
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text(x_tolerance=2, y_tolerance=2)
            if page_text:
                text_parts.append(page_text)
    return "\n".join(text_parts)


def extract_text_from_docx(content: bytes) -> str:
    doc = Document(io.BytesIO(content))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also extract table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())
    return "\n".join(paragraphs)
